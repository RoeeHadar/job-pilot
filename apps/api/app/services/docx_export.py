from __future__ import annotations

from io import BytesIO

from docx import Document


def markdown_to_docx_bytes(markdown: str, title: str = "Tailored CV") -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)
    for line in markdown.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line.strip())
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
